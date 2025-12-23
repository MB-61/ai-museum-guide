"""
DOCX → TXT Dönüştürücü
Tüm newDATA klasöründeki DOCX dosyalarını ted_museum/ klasörüne TXT olarak kaydeder.

Kullanım:
  python ingestion/convert_docx.py
"""
import os
import re
import json
from docx import Document

def slugify(text):
    """Dosya adından exhibit_id oluştur"""
    # Numarayı kaldır (örn: "1. " veya "10. ")
    text = re.sub(r'^\d+\.\s*', '', text)
    # Parantez içindeki yılı ayır
    year_match = re.search(r'\((\d{4})\)', text)
    year = year_match.group(1) if year_match else None
    # Parantezleri kaldır
    text = re.sub(r'\([^)]*\)', '', text)
    # Türkçe karakterleri dönüştür
    tr_map = str.maketrans('çğıöşüÇĞİÖŞÜ', 'cgiosuCGIOSU')
    text = text.translate(tr_map)
    # Sadece alfanumerik ve boşluk bırak
    text = re.sub(r'[^a-zA-Z0-9\s]', '', text)
    # Boşlukları alt çizgiye çevir ve küçült
    text = '_'.join(text.lower().split())
    # Yılı ekle
    if year:
        text = f"{text}_{year}"
    return text.strip('_')

def extract_text_from_docx(filepath):
    """DOCX dosyasından tüm metni çıkar"""
    doc = Document(filepath)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return '\n\n'.join(paragraphs)

def extract_title_from_docx(filepath):
    """İlk paragraftan başlığı al"""
    doc = Document(filepath)
    for p in doc.paragraphs:
        if p.text.strip():
            return p.text.strip()
    return os.path.basename(filepath)

def convert_all():
    source_dir = 'data/newDATA'
    target_dir = 'data/ted_museum'
    mapping_file = 'data/mappings/qr_to_exhibit.json'
    
    os.makedirs(target_dir, exist_ok=True)
    
    # Mevcut QR mapping'i yükle
    if os.path.exists(mapping_file):
        with open(mapping_file, 'r', encoding='utf-8') as f:
            qr_mapping = json.load(f)
    else:
        qr_mapping = {}
    
    files = sorted([f for f in os.listdir(source_dir) if f.endswith('.docx') and not f.startswith('~')])
    
    converted = []
    for i, filename in enumerate(files, 1):
        filepath = os.path.join(source_dir, filename)
        
        # exhibit_id oluştur
        base_name = os.path.splitext(filename)[0]
        exhibit_id = slugify(base_name)
        
        # Metni çıkar
        text = extract_text_from_docx(filepath)
        title = extract_title_from_docx(filepath)
        
        # TXT olarak kaydet
        txt_filename = f"{exhibit_id}.txt"
        txt_path = os.path.join(target_dir, txt_filename)
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write(text)
        
        # QR mapping güncelle
        qr_id = f"qr_{i:02d}"
        qr_mapping[qr_id] = exhibit_id
        
        converted.append({
            'qr_id': qr_id,
            'exhibit_id': exhibit_id,
            'title': title,
            'source': filename
        })
        print(f"[OK] {qr_id}: {exhibit_id}")
    
    # QR mapping kaydet
    with open(mapping_file, 'w', encoding='utf-8') as f:
        json.dump(qr_mapping, f, ensure_ascii=False, indent=2)
    
    print(f"\n[DONE] {len(converted)} dosya donusturuldu")
    print(f"[DIR] TXT dosyalari: {target_dir}/")
    print(f"[QR] QR Mapping: {mapping_file}")
    
    return converted

if __name__ == '__main__':
    convert_all()
